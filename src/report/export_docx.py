"""Export DOCX opzionale via python-docx."""

from __future__ import annotations

from pathlib import Path


class DocxExporter:
    """Esporta contenuto testuale in DOCX se python-docx e' disponibile."""

    def export(self, markdown_content: str, path: str | Path) -> Path:
        output_path = Path(path)
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dipendenza opzionale
            raise RuntimeError(
                "python-docx non disponibile. Installare con 'pip install python-docx'."
            ) from exc

        document = Document()
        for line in markdown_content.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.strip():
                document.add_paragraph(line)
        document.save(str(output_path))
        return output_path
